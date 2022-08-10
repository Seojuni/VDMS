from docx import Document
from docx.shared import Cm, Pt, RGBColor
import socket
import platform
from datetime import datetime
import mariadb
import db_conn
import matplotlib.pyplot as plt
from matplotlib import font_manager, rc
from docx2pdf import convert
import os
from PyPDF2 import PdfFileMerger


def create_pie(cur):
    select_query = 'select * from diagnostic_results'
    total_cnt = 0
    vuln = []
    pc_categorie1 = ['PC-01', 'PC-02', 'PC-15']
    pc_categorie2 = ['PC-03', 'PC-04', 'PC-05', 'PC-16', 'PC-17', 'PC-18']
    pc_categorie3 = ['PC-06', 'PC-07', 'PC-08']
    pc_categorie4 = ['PC-09', 'PC-10', 'PC-11', 'PC-12', 'PC-13', 'PC-14', 'PC-19']
    c1_cnt = 0
    c2_cnt = 0
    c3_cnt = 0
    c4_cnt = 0

    try:
        cur.execute(select_query)
    except mariadb.Error as e:
        print(f"Error: {e}")

    rows = cur.fetchall()
    for row in rows:
        if 'PC' in row[0]:
            total_cnt += 1
            if row[1] == '2':
                vuln.append(row[0])
                if row[0] in pc_categorie1:
                    c1_cnt += 1
                elif row[0] in pc_categorie2:
                    c2_cnt += 1
                elif row[0] in pc_categorie3:
                    c3_cnt += 1
                elif row[0] in pc_categorie4:
                    c4_cnt += 1
        if 'W' in row[0]:
            print('Server')

    font_path = "C:/Windows/Fonts/NGULIM.TTF"
    font = font_manager.FontProperties(fname=font_path).get_name()
    rc('font', family=font)

    vuln_per = '%.1f' % (len(vuln) / total_cnt * 100.0)
    good_per = '%.1f' % (100.0 - float(vuln_per.split('%')[0]))

    ratio = [good_per, vuln_per]
    labels = ['양호', '취약']
    colors = ['#8fd9b6', '#ff9999']
    explode = [0.05, 0]

    plt.figure(figsize=(5, 5), dpi=100)
    plt.pie(ratio, labels=labels, autopct='%.1f%%', colors=colors, explode=explode)
    plt.savefig('pie1.png')

    plt.clf()

    ratio = [c1_cnt, c2_cnt, c3_cnt, c4_cnt]
    labels = ['계정 관리', '서비스 관리', '패치 관리', '보안 관리']
    explode = [0.05, 0.05, 0.05, 0.05]

    plt.pie(ratio, labels=labels, autopct='%.1f%%', explode=explode)
    plt.savefig('pie2.png')

    return vuln


def create_cover(vuln_list):
    doc = Document()

    doc.add_heading('취약점 진단 결과 보고서', level=0)
    doc.add_heading('점검 대상', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = doc.styles['Table Grid']
    first_row = table.rows[0].cells
    first_row[0].text = '대상 이름'
    first_row[1].text = socket.gethostname()
    second_row = table.rows[1].cells
    second_row[0].text = '대상 운영체제'
    second_row[1].text = platform.system()
    third_row = table.rows[2].cells
    third_row[0].text = '점검 시각'
    third_row[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    doc.add_paragraph('')
    doc.add_paragraph('')

    doc.add_heading('점검 결과', level=1)
    p = doc.add_paragraph()
    p.add_run('  취약항목 : ').bold = True
    for vuln in vuln_list:
        if vuln == vuln_list[-1]:
            p.add_run(vuln)
        else:
            p.add_run(vuln + ', ')

    p = doc.add_paragraph()
    r = p.add_run()
    r.add_picture('pie1.png', width=Cm(7.5), height=Cm(7.5))
    r.add_picture('pie2.png', width=Cm(7.5), height=Cm(7.5))

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    para = doc.add_paragraph('해당 취약점 진단은 KISA 주요정보통신기반시설 기술적 취약점 분석ㆍ평가 방법 상세가이드를 기반하여 진단하였습니다.')
    para = doc.paragraphs[-1].runs
    doc.paragraphs[-1].runs[0].italic = True
    for run in para:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x78, 0x95, 0x3D)

    doc.save('report.docx')
    os.remove('pie1.png')
    os.remove('pie2.png')

    # docx to pdf
    inputFile = 'report.docx'
    outputFile = 'report.pdf'
    file = open(outputFile, 'w')
    file.close()

    convert(inputFile, outputFile)
    os.remove(inputFile)

    # merge report
    merger = PdfFileMerger()

    merger.append(outputFile)
    for vuln in vuln_list:
        merger.append("C:\\Users\\USER\\Desktop\\Report Source\\" + vuln +".pdf")

    merger.write("C:\\Users\\USER\\Desktop\\Report.pdf")
    merger.close()

    os.remove(outputFile)


if __name__ == "__main__":
    conn = db_conn.db_conn()
    cur = conn.cursor()

    vuln_list = create_pie(cur)
    create_cover(vuln_list)

    conn.close()
