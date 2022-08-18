from tkinter import *
import tkinter.messagebox
import customtkinter
import db_conn
import db_init
import mariadb

# dashboard
from PIL import Image, ImageTk
import create_dashboard_source as ds

# vulnerability diagnosis
from ttkwidgets import CheckboxTreeview
from tkinter import ttk
import datetime
import winreg as reg
import subprocess

# create report
from docx import Document
from docx.shared import Cm, Pt, RGBColor
import socket
import matplotlib.pyplot as plt
from matplotlib import font_manager, rc
from docx2pdf import convert
import os
from PyPDF2 import PdfFileMerger

# cve list
import webbrowser

# cve crawler
import threading
import schedule


customtkinter.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"


class App(customtkinter.CTk):
    WIDTH = 1280
    HEIGHT = 760

    def __init__(self):
        super().__init__()

        db_conn.db_conn()
        db_init.db_init()

        self.title("VDM.py")
        self.geometry(f"{App.WIDTH}x{App.HEIGHT}")
        self.protocol("WM_DELETE_WINDOW", self.on_closing)  # call .on_closing() when app gets closed

        # ============ create two frames ============
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.frame_left = customtkinter.CTkFrame(master=self,
                                                 width=180,
                                                 corner_radius=0)
        self.frame_left.grid(row=0, column=0, sticky="nswe")

        self.frame_right = customtkinter.CTkFrame(master=self)
        self.frame_right.grid(row=0, column=1, sticky="nswe", padx=20, pady=20)

        # ============ frame_left ============
        self.frame_left.grid_rowconfigure(0, minsize=10)
        self.frame_left.grid_rowconfigure(7, weight=1)
        self.frame_left.grid_rowconfigure(8, minsize=20)
        self.frame_left.grid_rowconfigure(11, minsize=10)

        self.label_1 = customtkinter.CTkLabel(master=self.frame_left,
                                              text="Vulnerability Diagnosis",
                                              text_font=("Roboto Medium", -16))
        self.label_1.grid(row=1, column=0, pady=0, padx=0)
        self.label_2 = customtkinter.CTkLabel(master=self.frame_left,
                                              text="and Management",
                                              text_font=("Roboto Medium", -16))
        self.label_2.grid(row=2, column=0, pady=10, padx=10)

        self.button_1 = customtkinter.CTkButton(master=self.frame_left,
                                                text="취약점 진단",
                                                command=self.one_button_event)
        self.button_1.grid(row=3, column=0, pady=10, padx=20)

        self.button_2 = customtkinter.CTkButton(master=self.frame_left,
                                                text="CVE List",
                                                command=self.two_button_event)
        self.button_2.grid(row=4, column=0, pady=10, padx=20)

        self.button_3 = customtkinter.CTkButton(master=self.frame_left,
                                                text="Vulnerability List",
                                                command=self.vuln_button_event)
        self.button_3.grid(row=5, column=0, pady=10, padx=20)

        self.button_4 = customtkinter.CTkButton(master=self.frame_left,
                                                text="Dashboard",
                                                command=self.three_button_event)
        self.button_4.grid(row=6, column=0, pady=10, padx=20)

        self.label_mode = customtkinter.CTkLabel(master=self.frame_left, text="Appearance Mode:")
        self.label_mode.grid(row=9, column=0, pady=0, padx=20, sticky="w")

        self.optionmenu_1 = customtkinter.CTkOptionMenu(master=self.frame_left,
                                                        values=["Light", "Dark", "System"],
                                                        command=self.change_appearance_mode)
        self.optionmenu_1.grid(row=10, column=0, pady=10, padx=20, sticky="w")

        # ============ frame_right ============
        self.frame_right.rowconfigure((0), weight=1)
        self.frame_right.columnconfigure((0), weight=1)

        self.frame_main = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_main.grid(row=0, column=0, sticky="nsew")
        self.frame_main.rowconfigure(2, weight=1)
        self.frame_main.columnconfigure(3, weight=1)

        # create dashboard source
        self.dashboard = ds.create_dashboard_source()
        self.dashboard.cwe_chart()
        self.dashboard.vuln_chart()
        self.vuln_total = self.dashboard.get_vuln_software_total()
        self.software_total = self.dashboard.get_installed_software_total()
        self.cve_total = self.dashboard.get_cve_total()

        self.dashboardImage1 = Image.open('resources\\chart1.png')
        photo = ImageTk.PhotoImage(self.dashboardImage1)
        self.dashboard1 = customtkinter.CTkLabel(master=self.frame_main,
                                                 width=660,
                                                 height=330,
                                                 image=photo,
                                                 corner_radius=6,
                                                 fg_color=("white", "white"),
                                                 justify=tkinter.LEFT)
        self.dashboard1.image = photo
        self.dashboard1.grid(column=0, row=0, columnspan=2, sticky="nwe", padx=10, pady=15)
        self.dashboard1_title = customtkinter.CTkLabel(master=self.frame_main,
                                                      text="CWE 유형 Top 10",
                                                      text_font=("", 13, "bold"),
                                                      text_color="gray38",
                                                      fg_color=("white", "white"))
        self.dashboard1_title.place(x=20, y=15)

        self.dashboardImage2 = Image.open('resources\\chart2.png')
        photo = ImageTk.PhotoImage(self.dashboardImage2)
        self.dashboard2 = customtkinter.CTkLabel(master=self.frame_main,
                                                 width=330,
                                                 height=330,
                                                 image=photo,
                                                 corner_radius=6,
                                                 fg_color=("white", "white"),
                                                 justify=tkinter.LEFT)
        self.dashboard2.image = photo
        self.dashboard2.grid(column=2, row=0, sticky="nwe", padx=10, pady=15)
        self.dashboard2_title = customtkinter.CTkLabel(master=self.frame_main,
                                                       text="취약점 점검 결과",
                                                       text_font=("", 13, "bold"),
                                                       text_color="gray38",
                                                       fg_color=("white", "white"))
        self.dashboard2_title.place(x=710, y=15)

        self.vuln_software_label = customtkinter.CTkLabel(master=self.frame_main,
                                                          text=self.vuln_total + "개",
                                                          width=330,
                                                          height=330,
                                                          text_font=("", 25,"bold"),
                                                          text_color="gray38",
                                                          corner_radius=6,
                                                          fg_color=("#E26365", "#E26365"),
                                                          justify=tkinter.LEFT)
        self.vuln_software_label.grid(column=0, row=1, sticky="nwe", padx=10, pady=5)
        self.vuln_software_title = customtkinter.CTkLabel(master=self.frame_main,
                                                          text="취약한 소프트웨어 버전",
                                                          text_font=("", 13, "bold"),
                                                          text_color="gray38",
                                                          fg_color=("#E26365", "#E26365"))
        self.vuln_software_title.place(x=25, y=380)

        self.installed_software_label = customtkinter.CTkLabel(master=self.frame_main,
                                                               text=self.software_total + "개",
                                                               width=330,
                                                               height=330,
                                                               text_font=("", 25, "bold"),
                                                               text_color="gray38",
                                                               corner_radius=6,
                                                               fg_color=("#00BAC7", "#00BAC7"),
                                                               justify=tkinter.LEFT)
        self.installed_software_label.grid(column=1, row=1, sticky="nwe", padx=5, pady=5)
        self.installed_software_title = customtkinter.CTkLabel(master=self.frame_main,
                                                               text="설치되어 있는 소프트웨어",
                                                               text_font=("", 13, "bold"),
                                                               text_color="gray38",
                                                               fg_color=("#00BAC7", "#00BAC7"))
        self.installed_software_title.place(x=370, y=380)

        self.cve_count_label = customtkinter.CTkLabel(master=self.frame_main,
                                                      text=self.cve_total + "개",
                                                      width=330,
                                                      height=330,
                                                      text_font=("", 25, "bold"),
                                                      text_color="gray38",
                                                      corner_radius=6,
                                                      fg_color=("#F0BF1D", "#F0BF1D"),
                                                      justify=tkinter.LEFT)
        self.cve_count_label.grid(column=2, row=1, sticky="nwe", padx=10, pady=5)
        self.cve_count_title = customtkinter.CTkLabel(master=self.frame_main,
                                                      text="검색 가능한 CVE Data",
                                                      text_font=("", 13, "bold"),
                                                      text_color="gray38",
                                                      fg_color=("#F0BF1D", "#F0BF1D"))
        self.cve_count_title.place(x=720, y=380)

        cve_crawler = threading.Thread(target=self.cve_crawler)
        cve_crawler.start()

        # set default values
        self.optionmenu_1.set("Dark")

    def cve_crawler(self):
        os.system("pythonw -u cve_crawler.py > log.log 2>&1")

        schedule.every(2).hour.do(self.cve_crawler)
        while True:
            schedule.run_pending()

    def vuln_comparison(self):
        os.system("pythonw -u vuln_comparison.py > log.log 2>&1")

    # go to vulnerability diagnosis
    def one_button_event(self):
        self.frame_one = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_one.grid(row=0, column=0, sticky="nsew")

        self.frame_one.rowconfigure(0, weight=1)
        self.frame_one.columnconfigure(0, weight=1)

        self.my_os = self.get_os()
        date = self.get_date()

        s = ttk.Style()
        s.theme_use("clam")
        s.configure("Treeview.Heading", font=(None, 14), rowheight=int(14 * 10))
        s.configure("Treeview", font=(None, 12), rowheight=int(12 * 5))

        tree = CheckboxTreeview(self.frame_one)
        tree.grid()

        tree["columns"] = ["one"]
        tree.heading("#0", text='취약점 진단(' + self.my_os + ')')
        tree.heading("#1", text=date)

        if self.my_os == "Windows PC":
            tree.insert('', tkinter.END, text='계정 관리', iid=1, open=False)
            tree.insert('', tkinter.END, text='서비스 관리', iid=2, open=False)
            tree.insert('', tkinter.END, text='패치 관리', iid=3, open=False)
            tree.insert('', tkinter.END, text='보안 관리', iid=4, open=False)

            acc = ['[PC-01] (상) 패스워드의 주기적 변경', '[PC-02] (상) 패스워드 정책이 해당 기관의 보안 정책에 적합하게 설정',
                   '[PC-15] (중) 복구 콘솔에서 자동 로그온을 금지하도록 설정']

            for i in range(len(acc)):
                code = acc[i].split(' ')[0]
                tree.insert('', tkinter.END, text=acc[i], iid=code, open=False)

            for i in range(len(acc)):
                code = acc[i].split(' ')[0]
                for j in range(len(acc)):
                    tree.move(code, 1, j)

            ser = ['[PC-03] (상) 공유 폴더 제거', '[PC-04] (상) 항목의 불필요한 서비스 제거',
                   '[PC-05] (상) Windows Messenger와 같은 상용 메신저의 사용 금지', '[PC-16] (중) 파일시스템이 NTFS 포맷으로 설정',
                   '[PC-17] (중) 대상 시스템이 Windows 서버를 제외한 다른 OS로 멀티 부팅이 가능하지 않도록 설정',
                   '[PC-18] (하) 브라우저 종료 시 임시 인터넷 파일 폴더의 내용을 삭제하도록 설정']

            for i in range(len(ser)):
                code = ser[i].split(' ')[0]
                tree.insert('', tkinter.END, text=ser[i], iid=code, open=False)

            for i in range(len(ser)):
                code = ser[i].split(' ')[0]
                for j in range(len(ser)):
                    tree.move(code, 2, j)

            patch = ['[PC-06] (상) HOT FIX 등 최신 보안패치 적용', '[PC-07] (상) 최신 서비스팩 적용',
                     '[PC-08] (상) MS-Office, 한글, 어도비 아크로뱃 등 응용 프로그램에 대한 최신 보안패치 및 벤더 권고사항 적용']

            for i in range(len(patch)):
                code = patch[i].split(' ')[0]
                tree.insert('', tkinter.END, text=patch[i], iid=code, open=False)

            for i in range(len(patch)):
                for j in range(len(patch)):
                    code = patch[i].split(' ')[0]
                    tree.move(code, 3, j)

            sec = ['[PC-09] (상) 바이러스 백신 프로그램 설치 및 주기적 업데이트', '[PC-10] (상) 바이러스 백신 프로그램에서 제공하는 실시간 감시 기능 활성화',
                   '[PC-11] (상) OS에서 제공하는 침입차단 기능 활성화', '[PC-12] (상) 화면보호기 대기 시간 설정 및 재시작 시 암호 보호 설정',
                   '[PC-13] (상) CD, DVD, USB 메모리 등과 같은 미디어의 자동실행 방지 등 이동식 미디어에 대한 보안대책 수립',
                   '[PC-14] (상) PC 내부의 미사용(3개월) ActiveX 제거', '[PC-19] (중) 원격 지원을 금지하도록 정책 설정']

            for i in range(len(sec)):
                code = sec[i].split(' ')[0]
                tree.insert('', tkinter.END, text=sec[i], iid=code, open=False)

            for i in range(len(sec)):
                code = sec[i].split(' ')[0]
                for j in range(len(sec)):
                    tree.move(code, 4, j)

        elif self.my_os == "Windows Server":
            tree.insert('', tkinter.END, text='계정 관리', iid=1, open=False)
            tree.insert('', tkinter.END, text='서비스 관리', iid=2, open=False)
            tree.insert('', tkinter.END, text='패치 관리', iid=3, open=False)
            tree.insert('', tkinter.END, text='로그 관리', iid=4, open=False)
            tree.insert('', tkinter.END, text='보안 관리', iid=5, open=False)
            tree.insert('', tkinter.END, text='DB 관리', iid=6, open=False)

            acc = ['[W-01] (상) Administrator 계정 이름 변경 또는 보안성 강화', '[W-02] (상) Guest 계정 비활성화', '[W-03] (상) 불필요한 계정 제거',
                   '[W-04] (상) 계정 잠금 임계값 설정', '[W-05] (상) 해독 가능한 암호화를 사용하여 암호 저장 해제', '[W-06] (상) 관리자 그룹에 최소한의 사용자 포함',
                   '[W-46] (중) Everyone 사용권한을 익명 사용자에 적용 해제', '[W-47] (중) 계정 잠금 기간 설정', '[W-48] (중) 패스워드 복잡성 설정',
                   '[W-49] (중) 패스워드 최소 암호 길이', '[W-50] (중) 패스워드 최대 사용 기간', '[W-51] (중) 패스워드 최소 사용 기간',
                   '[W-52] (중) 마지막 사용자 이름 표시 안함', '[W-53] (중) 로컬 로그온 허용', '[W-54] (중) 익명 SID/이름 변환 허용 해제',
                   '[W-55] (중) 최근 암호 기억', '[W-56] (중) 콘솔 로그온 시 로컬 계정에서 빈 암호 사용 제한', '[W-57] (중) 원격터미널 접속 가능한 사용자 그룹 제한']

            for i in range(len(acc)):
                code = acc[i].split(' ')[0]
                tree.insert('', tkinter.END, text=acc[i], iid=code, open=False)

            for i in range(len(acc)):
                code = acc[i].split(' ')[0]
                for j in range(len(acc)):
                    tree.move(code, 1, j)

            ser = ['[W-07] (상) 공유 권한 및 사용자 그룹 설정', '[W-08] (상) 하드디스크 기본 공유 제거', '[W-09] (상) 불필요한 서비스 제거',
                   '[W-10] (상) IIS 서비스 구동 점검', '[W-11] (상) IIS 디렉토리 리스팅 제거', '[W-12] (상) IIS CGI 실행 제한',
                   '[W-13] (상) IIS 상위 디렉토리 접근 금지', '[W-14] (상) IIS 불필요한 파일 제거', '[W-15] (상) IIS 웹프로세스 권한 제한',
                   '[W-16] (상) IIS 링크 사용 금지', '[W-17] (상) IIS 파일 업로드 및 다운로드 제한', '[W-18] (상) IIS DB 연결 취약점 점검',
                   '[W-19] (상) IIS 가상 디렉토리 삭제', '[W-20] (상) IIS 데이터파일 ACL 적용', '[W-21] (상) IIS 미사용 스크립트 매핑 제거',
                   '[W-22] (상) IIS Exec 명령어 쉘 호출 진단', '[W-23] (상) IIS WebDAV 비활성화', '[W-24] (상) NetBIOS 바인딩 서비스 구동 점검',
                   '[W-25] (상) FTP 서비스 구동 점검', '[W-26] (상) FTP 디렉토리 접근 권한 설정', '[W-27] (상) Anonymous FTP 금지',
                   '[W-28] (상) FTP 접근 제어 설정', '[W-29] (상) DNS Zone Transfer 설정', '[W-30] (상) RDS(Remonte Data Services) 제거',
                   '[W-31] (상) 최신 서비스팩 적용', '[W-58] (중) 터미널 서비스 암호화 수준 설정', '[W-59] (중) IIS 웹 서비스 정보 숨김',
                   '[W-60] (중) SNMP 서비스 구동 점검', '[W-61] (중) SNMP 서비스 커뮤니티스트링의 복잡성 설정', '[W-62] (중) SNMP Access control 설정',
                   '[W-63] (중) DNS 서비스 구동 점검', '[W-64] (하) HTTP/FTP/SMTP 배너 차단', '[W-65] (중) Telnet 보안 설정',
                   '[W-66] (중) 불필요한 ODBC/OLE=DB 데이터소스와 드라이브 제거', '[W-67] (중) 원격터미널 접속 타임아웃 설정',
                   '[W-68] (중) 예약된 작업에 의심스러운 명령이 등록되어 있는지 점검']

            for i in range(len(ser)):
                code = ser[i].split(' ')[0]
                tree.insert('', tkinter.END, text=ser[i], iid=code, open=False)

            for i in range(len(ser)):
                code = ser[i].split(' ')[0]
                for j in range(len(ser)):
                    tree.move(code, 2, j)

            patch = ['[W-32] (상) 최신 HOT FIX 적용', '[W-33] (상) 백신 프로그램 업데이트', '[W-69] (중) 정책에 따른 시스템 로깅설정']

            for i in range(len(patch)):
                code = patch[i].split(' ')[0]
                tree.insert('', tkinter.END, text=patch[i], iid=code, open=False)

            for i in range(len(patch)):
                for j in range(len(patch)):
                    code = patch[i].split(' ')[0]
                    tree.move(code, 3, j)

            log = ['[W-34] (상) 로그의 정기적 검토 및 보고', '[W-35] (상) 원격으로 엑세스 할 수 있는 레지스트리 경로',
                   '[W-70] (하) 이벤트 로그 관리 설정', '[W-71] (중) 원격에서 이벤트 로그파일 접근 차단']

            for i in range(len(log)):
                code = log[i].split(' ')[0]
                tree.insert('', tkinter.END, text=log[i], iid=code, open=False)

            for i in range(len(log)):
                for j in range(len(log)):
                    code = log[i].split(' ')[0]
                    tree.move(code, 4, j)

            sec = ['[W-36] (상) 백신 프로그램 설치', '[W-37] (상) SAM 파일 접근 통제 설정', '[W-38] (상) 화면보호기 설정',
                   '[W-39] (상) 로그온 하지않고 시스템 종료 허용 해제', '[W-40] (상) 원격 시스템에서 강제로 시스템 종료',
                   '[W-41] (상) 보안감사를 로그할 수 없는 경우 즉시 시스템 종료 해제', '[W-42] (상) SAM 계정과 공유의 익명 열거 허용 안함',
                   '[W-43] (상) Autologon 기능 제어', '[W-44] (상) 이동식 미디어 포맷 및 꺼내기 허용', '[W-45] (상) 디스크 볼륨 암호화 설정',
                   '[W-72] (중) Dos 공격 방어 레지스트리 설정', '[W-73] (중) 사용자가 프린터 드라이버를 설치할 수 없게 함',
                   '[W-74] (중) 세션 연결을 중단하기 전에 필요한 유휴시간', '[W-75] (하) 경고 메시지 설정', '[W-76] (중) 사용자별 홈 디렉토리 권한 설정',
                   '[W-77] (중) LAN Manager 인증 수준', '[W-78] (중) 보안 채널 데이터 디지털 암호화 또는 서명', '[W-79] (중) 파일 및 디렉토리 보호',
                   '[W-80] (중) 컴퓨터 계정 암호 최대 사용 기간', '[W-81] (중) 시작 프로그램 목록 분석']

            for i in range(len(sec)):
                code = sec[i].split(' ')[0]
                tree.insert('', tkinter.END, text=sec[i], iid=code, open=False)

            for i in range(len(sec)):
                code = sec[i].split(' ')[0]
                for j in range(len(sec)):
                    tree.move(code, 5, j)

            db = ['[W-82] (중) Windows 인증 모드 사용']

            for i in range(len(db)):
                code = db[i].split(' ')[0]
                tree.insert('', tkinter.END, text=db[i], iid=code, open=False)

            for i in range(len(db)):
                code = db[i].split(' ')[0]
                for j in range(len(db)):
                    tree.move(code, 6, j)

        tree.grid(row=0, column=0, sticky='nsew')

        self.button_start = customtkinter.CTkButton(master=self.frame_one,
                                                    text="Start",
                                                    command=lambda: self.vuln_diagnosis(tree))
        self.button_start.grid(row=1, column=0, pady=10, padx=20)


    # go to cve list
    def two_button_event(self):
        self.frame_two = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_two.grid(row=0, column=0, sticky="nsew")

        self.frame_two.rowconfigure(0, weight=1)
        self.frame_two.columnconfigure(0, weight=1)
        self.label_two = customtkinter.CTkLabel(master=self.frame_two,
                                                   text="CVE LIST",
                                                   height=20,
                                                   corner_radius=6,  # <- custom corner radius
                                                   fg_color=("white", "gray38"),  # <- custom tuple-color
                                                   justify=tkinter.LEFT)
        self.label_two.pack(fill="x", padx=15, pady=15)

        s = ttk.Style()
        s.theme_use("clam")
        s.configure("Treeview.Heading", font=(None, 14), rowheight=int(14 * 10))
        s.configure("Treeview", font=(None, 12), rowheight=int(12 * 5))

        self.treeview = tkinter.ttk.Treeview(self.frame_two,
                                        height=50,
                                        columns=["vuln_id", "pub_date", "last_mod_date", "description", "score", "software",
                                                 "start_version", "end_version", "cwe_id", "cwe_name"],
                                        displaycolumns=["vuln_id", "pub_date", "last_mod_date", "description", "score",
                                                        "software", "start_version", "end_version", "cwe_id", "cwe_name"]
                                        )
        ysb = ttk.Scrollbar(self.frame_two, orient="vertical", command=self.treeview.yview)
        ysb.pack(side='right', fill='y', pady=15)
        self.treeview.configure(yscrollcommand=ysb.set)
        self.treeview.pack(fill="x", padx=15, pady=15)

        self.treeview.column("#0", width=40)
        self.treeview.heading("#0", text="Num")
        self.treeview.column("vuln_id", width=150, anchor="center")
        self.treeview.heading("vuln_id", text="CVE Code", anchor="center")
        self.treeview.column("pub_date", width=150, anchor="center")
        self.treeview.heading("pub_date", text="Published Date", anchor="center")
        self.treeview.column("last_mod_date", width=150, anchor="center")
        self.treeview.heading("last_mod_date", text="Last Modified", anchor="center")
        self.treeview.column("description", width=150, anchor="center")
        self.treeview.heading("description", text="Description", anchor="center")
        self.treeview.column("score", width=100, anchor="center")
        self.treeview.heading("score", text="Score", anchor="center")
        self.treeview.column("software", width=150, anchor="center")
        self.treeview.heading("software", text="Software", anchor="center")
        self.treeview.column("start_version", width=150, anchor="center")
        self.treeview.heading("start_version", text="Start Version", anchor="center")
        self.treeview.column("end_version", width=150, anchor="center")
        self.treeview.heading("end_version", text="End Version", anchor="center")
        self.treeview.column("cwe_id", width=150, anchor="center")
        self.treeview.heading("cwe_id", text="CWE ID", anchor="center")
        self.treeview.column("cwe_name", width=150, anchor="center")
        self.treeview.heading("cwe_name", text="CWE Name", anchor="center")

        self.treeview.bind("<Double-1>", self.OnDoubleClick)

        conn = db_conn.db_conn()
        cur = conn.cursor()

        select_query = "select * from cve_list"

        try:
            cur.execute(select_query)
        except mariadb.Error as e:
            print(f"Error: {e}")

        resultset = cur.fetchall()

        for i, data in enumerate(resultset):
            self.treeview.insert("", tkinter.END, text=i+1, values=data, iid=data[0])

        conn.close()

    # view cve detail webpage
    def OnDoubleClick(self, event):
        item = self.treeview.selection()[0]
        webbrowser.open('https://nvd.nist.gov/vuln/detail/'+self.treeview.item(item, "values")[0])

    #go to Vulnerability List
    def vuln_button_event(self):
        self.frame_vuln = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_vuln.grid(row=0, column=0, sticky="nsew")

        self.frame_vuln.rowconfigure(0, weight=1)
        self.frame_vuln.columnconfigure(0, weight=1)
        self.label_vuln = customtkinter.CTkLabel(master=self.frame_vuln,
                                                text="Vulnerability LIST",
                                                height=20,
                                                corner_radius=6,  # <- custom corner radius
                                                fg_color=("white", "gray38"),  # <- custom tuple-color
                                                justify=tkinter.LEFT)
        self.label_vuln.pack(fill="x", padx=15, pady=15)

        s = ttk.Style()
        s.theme_use("clam")
        s.configure("Treeview.Heading", font=(None, 14), rowheight=int(14 * 10))
        s.configure("Treeview", font=(None, 12), rowheight=int(12 * 5))

        self.vuln_treeview = tkinter.ttk.Treeview(self.frame_vuln,
                                             height=50,
                                             columns=["name", "version", "vuln_id"],
                                             displaycolumns=["name", "version", "vuln_id"]
                                             )
        ysb = ttk.Scrollbar(self.frame_vuln, orient="vertical", command=self.vuln_treeview.yview)
        ysb.pack(side='right', fill='y', pady=15)
        self.vuln_treeview.configure(yscrollcommand=ysb.set)
        self.vuln_treeview.pack(fill="x", padx=15, pady=15)

        self.vuln_treeview.column("#0", width=40)
        self.vuln_treeview.heading("#0", text="Num")
        self.vuln_treeview.column("name", width=150, anchor="center")
        self.vuln_treeview.heading("name", text="Software Name", anchor="center")
        self.vuln_treeview.column("version", width=150, anchor="center")
        self.vuln_treeview.heading("version", text="Software Version", anchor="center")
        self.vuln_treeview.column("vuln_id", width=150, anchor="center")
        self.vuln_treeview.heading("vuln_id", text="CVE Code", anchor="center")

        self.vuln_treeview.bind("<Double-1>", self.vuln_OnDoubleClick)

        conn = db_conn.db_conn()
        cur = conn.cursor()

        select_query = "select * from vuln_list"

        try:
            cur.execute(select_query)
        except mariadb.Error as e:
            print(f"Error: {e}")

        resultset = cur.fetchall()

        for i, data in enumerate(resultset):
            self.vuln_treeview.insert("", tkinter.END, text=i + 1, values=data, iid=data[0])

        conn.close()

        # vuln_comparison = threading.Thread(target=self.vuln_comparison)
        # vuln_comparison.start()

    # view cve detail webpage
    def vuln_OnDoubleClick(self, event):
        item = self.vuln_treeview.selection()[0]
        webbrowser.open('https://nvd.nist.gov/vuln/detail/' + self.vuln_treeview.item(item, "values")[2])

    # go to Dashboard
    def three_button_event(self):
        self.frame_main = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_main.grid(row=0, column=0, sticky="nsew")
        self.frame_main.rowconfigure(2, weight=1)
        self.frame_main.columnconfigure(3, weight=1)

        self.dashboard = ds.create_dashboard_source()
        self.dashboard.cwe_chart()
        self.dashboard.vuln_chart()
        self.vuln_total = self.dashboard.get_vuln_software_total()
        self.software_total = self.dashboard.get_installed_software_total()
        self.cve_total = self.dashboard.get_cve_total()

        self.dashboardImage1 = Image.open('resources\\chart1.png')
        photo = ImageTk.PhotoImage(self.dashboardImage1)
        self.dashboard1 = customtkinter.CTkLabel(master=self.frame_main,
                                                 width=660,
                                                 height=330,
                                                 image=photo,
                                                 corner_radius=6,
                                                 fg_color=("white", "white"),
                                                 justify=tkinter.LEFT)
        self.dashboard1.image = photo
        self.dashboard1.grid(column=0, row=0, columnspan=2, sticky="nwe", padx=10, pady=15)
        self.dashboard1_title = customtkinter.CTkLabel(master=self.frame_main,
                                                       text="CWE 유형 Top 10",
                                                       text_font=("", 13, "bold"),
                                                       text_color="gray38",
                                                       fg_color=("white", "white"))
        self.dashboard1_title.place(x=20, y=15)

        self.dashboardImage2 = Image.open('resources\\chart2.png')
        photo = ImageTk.PhotoImage(self.dashboardImage2)
        self.dashboard2 = customtkinter.CTkLabel(master=self.frame_main,
                                                 width=330,
                                                 height=330,
                                                 image=photo,
                                                 corner_radius=6,
                                                 fg_color=("white", "white"),
                                                 justify=tkinter.LEFT)
        self.dashboard2.image = photo
        self.dashboard2.grid(column=2, row=0, sticky="nwe", padx=10, pady=15)
        self.dashboard2_title = customtkinter.CTkLabel(master=self.frame_main,
                                                       text="취약점 점검 결과",
                                                       text_font=("", 13, "bold"),
                                                       text_color="gray38",
                                                       fg_color=("white", "white"))
        self.dashboard2_title.place(x=710, y=15)

        self.vuln_software_label = customtkinter.CTkLabel(master=self.frame_main,
                                                          text=self.vuln_total + "개",
                                                          width=330,
                                                          height=330,
                                                          text_font=("", 25, "bold"),
                                                          text_color="gray38",
                                                          corner_radius=6,
                                                          fg_color=("#e21f26", "#e21f26"),
                                                          justify=tkinter.LEFT)
        self.vuln_software_label.grid(column=0, row=1, sticky="nwe", padx=10, pady=5)
        self.vuln_software_title = customtkinter.CTkLabel(master=self.frame_main,
                                                          text="취약한 소프트웨어 버전",
                                                          text_font=("", 13, "bold"),
                                                          text_color="gray38",
                                                          fg_color=("#e21f26", "#e21f26"))
        self.vuln_software_title.place(x=25, y=380)

        self.installed_software_label = customtkinter.CTkLabel(master=self.frame_main,
                                                               text=self.software_total + "개",
                                                               width=330,
                                                               height=330,
                                                               text_font=("", 25, "bold"),
                                                               text_color="gray38",
                                                               corner_radius=6,
                                                               fg_color=("#009aa5", "#009aa5"),
                                                               justify=tkinter.LEFT)
        self.installed_software_label.grid(column=1, row=1, sticky="nwe", padx=5, pady=5)
        self.installed_software_title = customtkinter.CTkLabel(master=self.frame_main,
                                                               text="설치되어 있는 소프트웨어",
                                                               text_font=("", 13, "bold"),
                                                               text_color="gray38",
                                                               fg_color=("#009aa5", "#009aa5"))
        self.installed_software_title.place(x=370, y=380)

        self.cve_count_label = customtkinter.CTkLabel(master=self.frame_main,
                                                      text=self.cve_total + "개",
                                                      width=330,
                                                      height=330,
                                                      text_font=("", 25, "bold"),
                                                      text_color="gray38",
                                                      corner_radius=6,
                                                      fg_color=("#ffcb1f", "#ffcb1f"),
                                                      justify=tkinter.LEFT)
        self.cve_count_label.grid(column=2, row=1, sticky="nwe", padx=10, pady=5)
        self.cve_count_title = customtkinter.CTkLabel(master=self.frame_main,
                                                      text="검색 가능한 CVE Data",
                                                      text_font=("", 13, "bold"),
                                                      text_color="gray38",
                                                      fg_color=("#ffcb1f", "#ffcb1f"))
        self.cve_count_title.place(x=720, y=380)

    def get_os(self):
        key = reg.HKEY_LOCAL_MACHINE
        key_value = "SOFTWARE\Microsoft\Windows NT\CurrentVersion"

        open = reg.OpenKey(key, key_value, 0, reg.KEY_READ)
        value, type = reg.QueryValueEx(open, "ProductName")

        reg.CloseKey(open)

        os = " ".join(str(value).split(' ')[0:2])

        if os == "Windows 10" or os == "Windows 11" or os == "Windows 8":
            return "Windows PC"
        else:
            return "Windows Server"

    def get_date(self):
        week = datetime.date.today().weekday()
        days = str(datetime.date.today().day)
        month = datetime.date.today().month
        year = str(datetime.date.today().year)

        weeks = ['Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat', 'Sun']
        months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
        date = weeks[week] + ", " + days + " " + months[month] + " " + year

        return date

    def vuln_diagnosis(self, tree):
        self.frame_one = customtkinter.CTkFrame(master=self.frame_right)
        self.frame_one.grid(row=0, column=0, sticky="nsew")

        self.frame_one.rowconfigure(0, weight=1)
        self.frame_one.columnconfigure(0, weight=1)

        selected = tree.get_checked()
        flg = 0
        self.digdate = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # textbox log
        self.scroll = tkinter.Scrollbar(self.frame_one, orient='vertical')
        self.lbox = tkinter.Listbox(self.frame_one, yscrollcommand=self.scroll.set, width=116)
        self.scroll.config(command=self.lbox.yview)
        self.lbox.grid(row=0, column=0, columnspan=5, sticky="nsew")
        self.append_log('Started diagnosing vulnerabilities.')

        # progress bar
        self.progressbar = customtkinter.CTkProgressBar(master=self.frame_one)
        self.progressbar.grid(row=1, column=0, sticky="ew", padx=15, pady=15)
        self.progressbar.set(0)

        self.button_save = customtkinter.CTkButton(master=self.frame_one,
                                                   text="Save",
                                                   command=self.create_report,
                                                   state="disabled")
        self.button_save.grid(row=2, column=0, pady=10, padx=20)

        # selected batch run
        for item in selected:
            self.append_log('Start diagnosing the ' + item[1:-1])

            subprocess.call(['C:\\Users\\USER\\Desktop\\vdm\\script\\' + self.my_os + '\\' + item[1:-1] + '.bat'])
            #print('C:\\Users\\USER\\Desktop\\vdm\\script\\' + self.my_os + '\\' + item[1:-1] + '.bat')

            self.progressbar.set((selected.index(item)+1) / len(selected) * 1.0)
            self.append_log(item[1:-1] + ' has been diagnosed.')

        flg = 1
        self.append_log('Vulnerability diagnosis is complete.')

        if flg == 1:
            self.button_save.configure(state="enabled")

    def append_log(self, msg):
        self.now = str(datetime.datetime.now())[0:-7]
        self.lbox.insert(tkinter.END, "[{}] {}".format(self.now, msg))
        self.lbox.update()
        self.lbox.see(tkinter.END)

    def create_report(self):
        conn = db_conn.db_conn()
        cur = conn.cursor()

        # save diagnostic result
        insert_query = 'insert into diagnostic_results(type_code, result) values (?, ?)'

        f = open("C:\\Users\\USER\\Desktop\\vdm\\script\\" + self.my_os + "\\result.txt", 'r')

        lines = f.readlines()
        for line in lines:
            type_code = line.strip().split(' ')[0]
            result = line.strip().split(' ')[1]

            try:
                cur.execute(insert_query, (type_code, result))
                conn.commit()
            except mariadb.Error as e:
                print(f"Error: {e}")

        f.close()
        os.remove("C:\\Users\\USER\\Desktop\\vdm\\script\\" + self.my_os + "\\result.txt")

        # create pie chart
        select_query = 'select * from diagnostic_results'
        total_cnt = 0
        vuln_list = []
        pc_categorie1 = ['PC-01', 'PC-02', 'PC-15']
        pc_categorie2 = ['PC-03', 'PC-04', 'PC-05', 'PC-16', 'PC-17', 'PC-18']
        pc_categorie3 = ['PC-06', 'PC-07', 'PC-08']
        pc_categorie4 = ['PC-09', 'PC-10', 'PC-11', 'PC-12', 'PC-13', 'PC-14', 'PC-19']
        server_categorie1 = ['W-01', 'W-02', 'W-03', 'W-04', 'W-05', 'W-06', 'W-46', 'W-47', 'W-48', 'W-49', 'W-50', 'W-51',
                             'W-52', 'W-53', 'W-54', 'W-55', 'W-56', 'W-57']
        server_categorie2 = ['W-07', 'W-08', 'W-09', 'W-10', 'W-11', 'W-12', 'W-13', 'W-14', 'W-15', 'W-16', 'W-17', 'W-18',
                             'W-19', 'W-20', 'W-21', 'W-22', 'W-23', 'W-24', 'W-25', 'W-26', 'W-27', 'W-28', 'W-29', 'W-30',
                             'W-31', 'W-58', 'W-59', 'W-60', 'W-61', 'W-62', 'W-63', 'W-64', 'W-65', 'W-66', 'W-67', 'W-68']
        server_categorie3 = ['W-32', 'W-33', 'W-69']
        server_categorie4 = ['W-34', 'W-35', 'W-70', 'W-71']
        server_categorie5 = ['W-36', 'W-37', 'W-38', 'W-39', 'W-40', 'W-41', 'W-42', 'W-43', 'W-44', 'W-45', 'W-72', 'W-73',
                             'W-74', 'W-75', 'W-76', 'W-77', 'W-78', 'W-79', 'W-80', 'W-81']
        server_categorie6 = ['W-82']
        c1_cnt = 0
        c2_cnt = 0
        c3_cnt = 0
        c4_cnt = 0
        c5_cnt = 0
        c6_cnt = 0

        try:
            cur.execute(select_query)
        except mariadb.Error as e:
            print(f"Error: {e}")

        rows = cur.fetchall()
        for row in rows:
            if 'PC' in row[0]:
                total_cnt += 1
                if row[1] == '2':
                    vuln_list.append(row[0])
                    if row[0] in pc_categorie1:
                        c1_cnt += 1
                    elif row[0] in pc_categorie2:
                        c2_cnt += 1
                    elif row[0] in pc_categorie3:
                        c3_cnt += 1
                    elif row[0] in pc_categorie4:
                        c4_cnt += 1
            if 'W' in row[0]:
                total_cnt += 1
                if row[1] == '2':
                    vuln_list.append(row[0])
                    if row[0] in server_categorie1:
                        c1_cnt += 1
                    elif row[0] in server_categorie2:
                        c2_cnt += 1
                    elif row[0] in server_categorie3:
                        c3_cnt += 1
                    elif row[0] in server_categorie4:
                        c4_cnt += 1
                    elif row[0] in server_categorie5:
                        c5_cnt += 1
                    elif row[0] in server_categorie6:
                        c6_cnt += 1

        font_path = "C:/Windows/Fonts/NGULIM.TTF"
        font = font_manager.FontProperties(fname=font_path).get_name()
        rc('font', family=font)

        vuln_per = '%.1f' % (len(vuln_list) / total_cnt * 100.0)
        good_per = '%.1f' % (100.0 - float(vuln_per.split('%')[0]))

        ratio = [good_per, vuln_per]
        labels = ['양호', '취약']
        colors = ['#8fd9b6', '#ff9999']
        explode = [0.05, 0]

        plt.figure(figsize=(5, 5), dpi=100)
        plt.pie(ratio, labels=labels, autopct='%.1f%%', colors=colors, explode=explode)
        plt.savefig('pie1.png')

        plt.clf()

        if self.my_os == "Windows PC":
            ratio = [c1_cnt, c2_cnt, c3_cnt, c4_cnt]
            labels = ['계정 관리', '서비스 관리', '패치 관리', '보안 관리']
            explode = [0.05, 0.05, 0.05, 0.05]

            plt.pie(ratio, labels=labels, autopct='%.1f%%', explode=explode)
            plt.savefig('pie2.png')
        elif self.my_os == "Windows Server":
            ratio = [c1_cnt, c2_cnt, c3_cnt, c4_cnt, c5_cnt, c6_cnt]
            labels = ['계정 관리', '서비스 관리', '패치 관리', '로그 관리', '보안 관리', 'DB 관리']
            explode = [0.05, 0.05, 0.05, 0.05, 0.05, 0.05]

            plt.pie(ratio, labels=labels, autopct='%.1f%%', explode=explode)
            plt.savefig('pie2.png')

        # create report cover
        doc = Document()

        doc.add_heading('취약점 진단 결과 보고서', level=0)
        doc.add_heading('점검 대상', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = doc.styles['Table Grid']
        first_row = table.rows[0].cells
        first_row[0].text = '대상 이름'
        first_row[1].text = socket.gethostname()
        second_row = table.rows[1].cells
        second_row[0].text = '대상 운영체제'
        second_row[1].text = self.my_os
        third_row = table.rows[2].cells
        third_row[0].text = '점검 시각'
        third_row[1].text = self.digdate

        doc.add_paragraph('')
        doc.add_paragraph('')

        doc.add_heading('점검 결과', level=1)
        p = doc.add_paragraph()
        p.add_run('  취약항목 : ').bold = True
        for vuln in vuln_list:
            if vuln == vuln_list[-1]:
                p.add_run(vuln)
            else:
                p.add_run(vuln + ', ')

        p = doc.add_paragraph()
        r = p.add_run()
        r.add_picture('pie1.png', width=Cm(7.5), height=Cm(7.5))
        r.add_picture('pie2.png', width=Cm(7.5), height=Cm(7.5))

        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')
        doc.add_paragraph('')

        para = doc.add_paragraph('해당 취약점 진단은 KISA 주요정보통신기반시설 기술적 취약점 분석ㆍ평가 방법 상세가이드를 기반하여 진단하였습니다.')
        para = doc.paragraphs[-1].runs
        doc.paragraphs[-1].runs[0].italic = True
        for run in para:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x78, 0x95, 0x3D)

        doc.save('report.docx')
        os.remove('pie1.png')
        os.remove('pie2.png')

        # report cover docx to pdf
        inputFile = 'report.docx'
        outputFile = 'report.pdf'
        file = open(outputFile, 'w')
        file.close()

        convert(inputFile, outputFile)
        os.remove(inputFile)

        # merge reports
        merger = PdfFileMerger()

        merger.append(outputFile)
        for vuln in vuln_list:
            merger.append("C:\\Users\\USER\\Desktop\\vdm\\Report Source\\" + self.my_os + "\\" + vuln + ".pdf")

        merger.write("C:\\Users\\USER\\Desktop\\Report.pdf")
        merger.close()

        os.remove(outputFile)

        tkinter.messagebox.showinfo("Save", "보고서가 저장되었습니다.")

    def change_appearance_mode(self, new_appearance_mode):
        customtkinter.set_appearance_mode(new_appearance_mode)

    def on_closing(self, event=0):
        self.destroy()


if __name__ == "__main__":
    app = App()
    app.mainloop()
